from copy import deepcopy
from io import BytesIO
from pathlib import Path
import datetime as dt
import os
import shutil
import subprocess
import tempfile

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from scoring import (
    FREE_DEBATE_CRITERIA,
    SPEECH_CRITERIA,
    free_debate_col,
    speech_col,
)


TEMPLATE_PATH = Path(__file__).resolve().parent / "assets" / "pdf_templates" / "score_sheet_template.docx"
FONT_NAME = "Noto Sans CJK TC"
TABLES_PER_PAGE = 8


def _is_blank(value):
    if value is None:
        return True
    try:
        return bool(pd.isna(value))
    except Exception:
        return False


def _get(source, key, default=""):
    try:
        value = source.get(key, default)
    except AttributeError:
        try:
            value = source[key]
        except Exception:
            value = default
    return default if _is_blank(value) else value


def _num(value, default=0):
    if _is_blank(value):
        return default
    try:
        return int(round(float(value)))
    except (TypeError, ValueError):
        return default


def _format_date(value):
    if _is_blank(value):
        return ""
    try:
        parsed = pd.to_datetime(value, errors="coerce")
    except Exception:
        parsed = None
    if parsed is None or pd.isna(parsed):
        return str(value)
    return f"{parsed.year} 年 {parsed.month} 月 {parsed.day} 日"


def _format_time(value):
    if _is_blank(value):
        return ""
    if isinstance(value, (dt.datetime, dt.time)):
        return value.strftime("%H:%M")
    text = str(value).strip()
    if len(text) >= 5 and text[2:3] == ":":
        return text[:5]
    try:
        parsed = pd.to_datetime(text, errors="coerce")
    except Exception:
        parsed = None
    if parsed is not None and not pd.isna(parsed):
        return parsed.strftime("%H:%M")
    return text


def _as_df(value):
    return value if isinstance(value, pd.DataFrame) else pd.DataFrame()


def _row_value(df, row_index, column, default=""):
    if df.empty or column not in df.columns or row_index >= len(df.index):
        return default
    return _get(df.iloc[row_index], column, default)


def _set_run_font(run, size=10, bold=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for name in ("eastAsia", "ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{name}"), FONT_NAME)


def _set_cell_text(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(str(text))
    _set_run_font(run, size=size, bold=bold)


def _weighted_speech_scores(df, row_index):
    scores = []
    for criterion in SPEECH_CRITERIA:
        raw_score = _num(_row_value(df, row_index, speech_col(criterion), 0))
        scores.append(raw_score * criterion["weight"])
    return scores


def _individual_scores(side_data):
    scores = side_data.get("ind_scores")
    if isinstance(scores, list) and len(scores) >= 4:
        return [_num(score) for score in scores[:4]]

    df = _as_df(side_data.get("raw_df_a"))
    return [sum(_weighted_speech_scores(df, i)) for i in range(4)]


def _build_ranks(pro_data, con_data):
    all_scores = _individual_scores(pro_data) + _individual_scores(con_data)
    if len(all_scores) != 8:
        ranks = [""] * 8
    else:
        ranks = pd.Series(all_scores).rank(ascending=False, method="min").astype(int).tolist()
    return {
        "正方": ranks[:4],
        "反方": ranks[4:],
    }


def _winner_label(judge_record):
    pro_total = _num(_get(judge_record, "pro_total_score"))
    con_total = _num(_get(judge_record, "con_total_score"))
    if pro_total > con_total:
        return "正方"
    if con_total > pro_total:
        return "反方"
    return "平局"


def _duplicate_template_page(doc):
    body = doc._body._element
    original_elements = [
        deepcopy(element)
        for element in list(body)
        if element.tag != qn("w:sectPr")
    ]

    page_break = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    page_break.append(run)

    insert_at = len(body) - 1 if body[-1].tag == qn("w:sectPr") else len(body)
    body.insert(insert_at, page_break)
    insert_at += 1
    for element in original_elements:
        body.insert(insert_at, element)
        insert_at += 1


def _fill_meta_table(table, match_info, side_label, team_name):
    _set_cell_text(table.cell(0, 1), _format_date(_get(match_info, "match_date")), align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(table.cell(0, 3), _format_time(_get(match_info, "match_time")), align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(table.cell(1, 1), f"{side_label}：{team_name}", size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(table.cell(1, 3), _get(match_info, "match_id"), size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(table.cell(2, 1), _get(match_info, "topic_text"), size=9, align=WD_ALIGN_PARAGRAPH.LEFT)


def _fill_speech_table(table, side_data, ranks):
    df = _as_df(side_data.get("raw_df_a"))
    for row_index in range(4):
        table_row = row_index + 1
        name = _row_value(df, row_index, "姓名", "")
        weighted_scores = _weighted_speech_scores(df, row_index)
        values = [name, *weighted_scores, sum(weighted_scores), ranks[row_index] if row_index < len(ranks) else ""]
        for col_index, value in enumerate(values, start=1):
            _set_cell_text(table.cell(table_row, col_index), value, size=10)


def _fill_free_debate_table(table, side_data):
    df = _as_df(side_data.get("raw_df_b"))
    total = 0
    for col_index, criterion in enumerate(FREE_DEBATE_CRITERIA):
        score = _num(_row_value(df, 0, free_debate_col(criterion), 0))
        total += score
        _set_cell_text(table.cell(1, col_index), score, size=10)
    _set_cell_text(table.cell(1, 5), total, size=10)


def _fill_deduction_table(table, side_data):
    for col_index in range(4):
        _set_cell_text(table.cell(1, col_index), "", size=10)
    _set_cell_text(table.cell(1, 4), _num(side_data.get("deduction")), size=10)


def _fill_footer_tables(totals_table, winner_table, signature_table, side_data, judge_record):
    total_a = _num(side_data.get("total_a"))
    total_b = _num(side_data.get("total_b"))
    deduction = _num(side_data.get("deduction"))
    coherence = _num(side_data.get("coherence"))
    subtotal = total_a + total_b - deduction
    final_total = _num(side_data.get("final_total"))

    _set_cell_text(totals_table.cell(0, 1), subtotal, size=10)
    _set_cell_text(totals_table.cell(0, 3), coherence, size=10)
    _set_cell_text(totals_table.cell(0, 5), final_total, size=10)
    _set_cell_text(winner_table.cell(0, 0), f"勝方：{_winner_label(judge_record)}", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(signature_table.cell(0, 1), _get(judge_record, "judge_name"), size=10, align=WD_ALIGN_PARAGRAPH.LEFT)


def _fill_score_sheet_page(tables, match_info, judge_record, side_label, team_name, side_data, ranks):
    _fill_meta_table(tables[1], match_info, side_label, team_name)
    _fill_speech_table(tables[2], side_data, ranks)
    _fill_free_debate_table(tables[3], side_data)
    _fill_deduction_table(tables[4], side_data)
    _fill_footer_tables(tables[5], tables[6], tables[7], side_data, judge_record)


def _build_score_sheet_docx(match_info, judge_record, pro_data, con_data):
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"DOCX template not found: {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))
    if len(doc.tables) != TABLES_PER_PAGE:
        raise ValueError(f"DOCX template should contain {TABLES_PER_PAGE} tables, found {len(doc.tables)}.")

    _duplicate_template_page(doc)
    if len(doc.tables) != TABLES_PER_PAGE * 2:
        raise ValueError("Failed to duplicate DOCX score sheet template.")

    ranks = _build_ranks(pro_data, con_data)
    pages = [
        (doc.tables[:TABLES_PER_PAGE], "正方", _get(judge_record, "pro_team"), pro_data),
        (doc.tables[TABLES_PER_PAGE:TABLES_PER_PAGE * 2], "反方", _get(judge_record, "con_team"), con_data),
    ]
    for tables, side_label, team_name, side_data in pages:
        _fill_score_sheet_page(tables, match_info, judge_record, side_label, team_name, side_data, ranks[side_label])

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _find_soffice():
    soffice = shutil.which("soffice")
    if soffice:
        return soffice

    mac_soffice = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if Path(mac_soffice).exists():
        return mac_soffice
    return None


def _convert_docx_to_pdf(docx_bytes):
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError("PDF 轉換工具未安裝：請在部署環境安裝 LibreOffice（soffice）。")

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        docx_path = tmp_path / "score_sheet.docx"
        pdf_path = tmp_path / "score_sheet.pdf"
        profile_dir = tmp_path / "lo_profile"
        docx_path.write_bytes(docx_bytes)

        env = os.environ.copy()
        env["HOME"] = str(tmp_path)
        cmd = [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            f"-env:UserInstallation=file://{profile_dir}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(tmp_path),
            str(docx_path),
        ]
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, env=env, timeout=60)
        if result.returncode != 0 or not pdf_path.exists():
            message = result.stderr.strip() or result.stdout.strip() or "LibreOffice PDF conversion failed."
            raise RuntimeError(f"PDF 轉換失敗：{message}")
        return pdf_path.read_bytes()


def build_score_sheet_pdf(match_info, judge_record, pro_data, con_data):
    docx_bytes = _build_score_sheet_docx(match_info, judge_record, pro_data, con_data)
    return _convert_docx_to_pdf(docx_bytes)
